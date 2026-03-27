# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 페이지 여백
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── 색상
BLUE  = RGBColor(0x1A, 0x6F, 0xD4)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GRAY  = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT = 'Malgun Gothic'

# ── run에 한글 폰트 적용
def kr(run, size=10.5, bold=False, color=None, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    run.font.size  = Pt(size)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color

# ── paragraph 기본 간격
def spacing(p, before=0, after=4):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

# ── 단락 하단 파란 테두리
def blue_border(p):
    pPr   = p._p.get_or_add_pPr()
    pBdr  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '8')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1A6FD4')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── 셀 배경색
def cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ────────────────────────────────────────────────
# 헬퍼: 제목1
def h1(text):
    p = doc.add_paragraph()
    spacing(p, before=18, after=6)
    blue_border(p)
    run = p.add_run(text)
    kr(run, size=16, bold=True, color=BLUE)

# 헬퍼: 제목2
def h2(text):
    p = doc.add_paragraph()
    spacing(p, before=12, after=4)
    run = p.add_run(text)
    kr(run, size=13, bold=True, color=DARK)

# 헬퍼: 제목3
def h3(text):
    p = doc.add_paragraph()
    spacing(p, before=8, after=2)
    run = p.add_run(text)
    kr(run, size=11, bold=True, color=BLUE)

# 헬퍼: 본문
def body(text, indent=False):
    p = doc.add_paragraph()
    spacing(p, after=4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    kr(run, size=10.5)

# 헬퍼: 글머리 기호
def bul(text, level=0):
    p = doc.add_paragraph()
    spacing(p, after=3)
    p.paragraph_format.left_indent  = Cm(0.8 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    run = p.add_run(f'• {text}')
    kr(run, size=10.5)

# 헬퍼: 노트 박스 (파란 배경)
def note(text):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    cell_bg(cell, 'EFF6FF')
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run(text)
    kr(run, size=10, color=DARK)
    doc.add_paragraph()

# 헬퍼: 표
def table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style      = 'Table Grid'
    tbl.alignment  = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell_bg(cell, '1A6FD4')
        cell.paragraphs[0].clear()
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        kr(run, size=10, bold=True, color=WHITE)

    # 데이터 행
    for ri, row in enumerate(rows):
        bg = 'FFFFFF' if ri % 2 == 0 else 'F0F7FF'
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            p   = cell.paragraphs[0]
            run = p.add_run(str(val))
            kr(run, size=10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()

# ════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(72)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('대학알리미 추천시스템')
kr(run, size=28, bold=True, color=BLUE)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('사용자 매뉴얼')
kr(run2, size=20, bold=True, color=DARK)

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Phase 1~2 데모 버전')
kr(run3, size=13, color=GRAY)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('버전: v1.0    작성일: 2026년 3월 27일')
kr(run4, size=11, color=GRAY)

doc.add_page_break()

# ════════════════════════════════════════════════
# 1. 개요
# ════════════════════════════════════════════════
h1('1. 개요')

h2('1.1 시스템 소개')
body('대학알리미 추천시스템(Phase 1~2 데모)는 대학알리미 포털(academyinfo.go.kr) 이용자의 행동 데이터를 분석하여 맞춤형 대학·지표를 추천하는 시스템의 핵심 기반 모듈을 구현한 데모입니다.')
body('이 데모에서는 ① 샘플 데이터 구축(데이터 수집 단계)과 ② 행동 분석 엔진(핵심 모듈)의 동작을 직접 확인할 수 있습니다.')

h2('1.2 주요 기능')
bul('Session ID 기반 사용자 식별 (UUID v4 + SHA-256 Fingerprint)')
bul('8종 행동 이벤트 자동 수집 (검색, 페이지 조회, 지표 확인, 필터, 비교 등)')
bul('Rule-based 의도 패턴 실시간 분류 (5종)')
bul('Session Context Document 자동 생성 (RAG + GPT-Oss 20B 프롬프트 입력용)')
bul('5개 시나리오 Replay 기능 (의도 패턴 분류 검증용)')

h2('1.3 접속 방법')
body('아래 URL로 접속합니다 (별도 설치 불필요):')
note('https://usingsys-hub.github.io/HEK/')

doc.add_page_break()

# ════════════════════════════════════════════════
# 2. 화면 구성
# ════════════════════════════════════════════════
h1('2. 화면 구성')
body('화면은 좌측 메인 영역과 우측 행동 분석 대시보드, 두 영역으로 나뉩니다.')

table(
    ['영역', '위치', '주요 내용'],
    [
        ['메인 영역', '좌측', '대학 검색, 지표 조회, Session Context Document 확인'],
        ['행동 분석 대시보드', '우측 (검은 배경)', '실시간 이벤트 로그, 의도 패턴, 시나리오 Replay'],
    ],
    [4, 3, 8.5]
)

h2('2.1 메인 영역 탭')
table(
    ['탭', '설명'],
    [
        ['대학 검색', '대학명·지표명 검색, 지역·유형 필터, 대학 카드 선택, 지표 테이블 조회'],
        ['Context Document', '현재 세션의 Session Context Document(RAG 입력용 자연어 요약) 확인'],
    ],
    [4, 11.5]
)

h2('2.2 행동 분석 대시보드 (우측)')
bul('세션 ID: 현재 세션의 UUID 표시')
bul('이벤트 수: 누적된 행동 이벤트 총 개수')
bul('감지된 의도 패턴: 현재 분석된 패턴을 색상 배지로 표시')
bul('시나리오 Replay: 5개 버튼으로 테스트 시나리오 즉시 주입')
bul('이벤트 로그: 최신순으로 이벤트 타입, 시각, 메타데이터 표시')
bul('세션 초기화: 버튼 클릭 시 새 Session ID 발급 및 이벤트 초기화')

doc.add_page_break()

# ════════════════════════════════════════════════
# 3. 사용 방법
# ════════════════════════════════════════════════
h1('3. 사용 방법')

h2('3.1 대학 검색')
bul('검색창에 대학명(예: 서울대학교) 또는 지표명(예: 장학금)을 입력합니다.')
bul('검색 버튼 클릭 또는 Enter 입력 시 SEARCH 이벤트가 기록됩니다.')
bul('검색 키워드가 지표명 목록에 포함되면 자동으로 지표명 검색으로 분류됩니다.')

h2('3.2 지역·유형 필터')
bul('지역 드롭다운: 특정 지역 선택 시 FILTER_APPLY(지역) 이벤트가 기록됩니다.')
bul('유형 드롭다운: 국립/사립 선택 시 FILTER_APPLY(대학유형) 이벤트가 기록됩니다.')
note('지역 필터를 세션 초반에 적용하면 의도 패턴이 REGION_FIRST로 분류될 수 있습니다.')

h2('3.3 대학 카드 선택 및 지표 조회')
bul('대학 카드 클릭 시 PAGE_VIEW(대학소개) 이벤트가 기록되고 하단에 지표 테이블이 펼쳐집니다.')
bul('지표 테이블의 각 행을 클릭하면 INDICATOR_VIEW 이벤트가 기록됩니다.')
bul('여러 대학 카드를 연속 클릭하면 COMPARE_MODE 패턴으로 분류됩니다.')
note('지표 행에 마우스를 올리면 배경이 파란색으로 변합니다. 클릭하면 이벤트가 기록됩니다.')

h2('3.4 Context Document 탭 확인')
bul('상단 [Context Document] 탭을 클릭합니다.')
bul('자연어 요약, 의도 패턴, 조회 대학, 관심 지표, 적용 필터, 체류 시간이 표시됩니다.')
bul('[Raw JSON 보기]를 펼치면 실제 JSON 구조를 확인할 수 있습니다.')

h2('3.5 시나리오 Replay')
body('우측 대시보드의 [시나리오 Replay] 버튼을 사용하면 미리 준비된 5개 시나리오를 즉시 주입하여 의도 패턴 분류 결과를 테스트할 수 있습니다.')

table(
    ['버튼', '시나리오 내용', '예상 패턴'],
    [
        ['UNIV_FIRST',      '"서울대학교" 검색 → 대학소개 조회 → 장학금·교육비·취업률 지표 확인',          'UNIV_FIRST'],
        ['INDICATOR_FIRST', '"재학생 장학금" 검색 → KAIST·POSTECH·성균관대 장학금 지표 비교 → KAIST 소개', 'INDICATOR_FIRST'],
        ['REGION_FIRST',    '부산 지역 필터 → 부산대·동아대 대학소개 → 취업률·전공취업률 지표 확인',       'REGION_FIRST'],
        ['COMPARE_MODE',    '서울대·연세대·고려대 순서로 페이지 조회 + COMPARE 이벤트 발생',               'COMPARE_MODE'],
        ['EXPLORATORY',     '홈 → 통합 검색 → 다양한 지표·대학 혼재 탐색 (패턴 없음)',                    'EXPLORATORY'],
    ],
    [3, 9, 3.5]
)

h2('3.6 세션 초기화')
bul('우측 대시보드 우상단의 [세션 초기화] 버튼을 클릭합니다.')
bul('현재 세션이 삭제되고 새 Session ID가 발급됩니다.')
bul('이벤트 로그와 의도 패턴이 초기화됩니다.')

doc.add_page_break()

# ════════════════════════════════════════════════
# 4. 이벤트 및 의도 패턴 상세
# ════════════════════════════════════════════════
h1('4. 이벤트 및 의도 패턴 상세')

h2('4.1 수집되는 이벤트 종류')
table(
    ['이벤트 타입', '발생 조건', '주요 메타데이터'],
    [
        ['SEARCH',         '검색 버튼 클릭',          'search_keyword, search_type, result_count'],
        ['PAGE_VIEW',      '대학 카드 클릭',           'page_url, page_type, university_id'],
        ['INDICATOR_VIEW', '지표 테이블 행 클릭',      'indicator_id, indicator_name, university_id'],
        ['FILTER_APPLY',   '지역/유형 드롭다운 변경',  'filter_type, filter_value'],
        ['COMPARE',        '시나리오 Replay(비교)',     'compared_universities, compared_indicators'],
        ['DWELL',          '시나리오 Replay(체류)',     'page_url, dwell_time_seconds, scroll_depth_percent'],
        ['DOWNLOAD',       'API 직접 호출',             'download_type, file_format'],
        ['CLICK',          'API 직접 호출',             'clicked_item, recommendation_id'],
    ],
    [3.5, 4.5, 7.5]
)

h2('4.2 의도 패턴 분류 규칙')
body('설계서 5.2 기준의 Rule-based 분류 규칙입니다. COMPARE_MODE가 최우선 규칙입니다.')

table(
    ['패턴', '분류 기준', '배지 색상'],
    [
        ['UNIV_FIRST',      '첫 3개 이벤트에 대학명 검색/대학소개 조회 포함 + 이후 지표 확인',  '파란색'],
        ['INDICATOR_FIRST', '첫 3개 이벤트에 지표명 검색/지표 조회 포함 + 이후 대학소개 이동', '초록색'],
        ['REGION_FIRST',    '이벤트 앞 절반에 지역 필터 적용 포함',                              '주황색'],
        ['COMPARE_MODE',    'COMPARE 이벤트 발생 OR 2개 이상 대학 PAGE_VIEW (최우선 규칙)',      '보라색'],
        ['EXPLORATORY',     '위 패턴 미해당 + 3가지 이상 이벤트 타입 혼재',                     '빨간색'],
        ['UNKNOWN',         '이벤트 3개 미만 또는 패턴 미확인',                                  '회색'],
    ],
    [3.5, 8, 2]
)

doc.add_page_break()

# ════════════════════════════════════════════════
# 5. 11개 주요 지표
# ════════════════════════════════════════════════
h1('5. 11개 주요 지표')
body('대학알리미 설계서 기준 11개 주요 지표입니다. 지표 테이블에서 각 항목을 클릭하면 INDICATOR_VIEW 이벤트로 기록됩니다.')

table(
    ['번호', '지표명', '단위', '설명'],
    [
        ['1',  '재학생1인당장학금',      '원',  '재학생 1인당 지급된 장학금 총액'],
        ['2',  '전임교원확보율',          '%',   '정원 대비 전임교원 확보 비율'],
        ['3',  '학생1인당교육비',         '원',  '학생 1인당 투입된 교육비'],
        ['4',  '학생1인당도서수',         '권',  '재학생 1인당 보유 도서 수'],
        ['5',  '외국인유학생1인당지원금', '원',  '외국인 유학생 1인당 지원 금액'],
        ['6',  '전임교원강의담당비율',    '%',   '전체 강의 중 전임교원 담당 비율'],
        ['7',  '단과대학생수',            '명',  '단과대학 재학생 수'],
        ['8',  '졸업후취업비율',          '%',   '졸업 후 취업한 비율'],
        ['9',  '전공관련취업자비율',      '%',   '취업자 중 전공 관련 직종 취업 비율'],
        ['10', '전공관련1인당연구비',     '원',  '전임교원 1인당 전공 관련 연구비'],
        ['11', '법인전입금',              '원',  '법인이 학교에 전입한 금액'],
    ],
    [1, 5, 1.5, 8]
)

# ════════════════════════════════════════════════
# 6. 샘플 대학 목록
# ════════════════════════════════════════════════
h1('6. 샘플 대학 목록 (20개)')
body('다음 20개 대학의 샘플 데이터가 포함되어 있습니다.')

table(
    ['ID', '대학명', '지역', '유형'],
    [
        ['UNIV_001', '서울대학교',     '서울', '국립'],
        ['UNIV_002', '연세대학교',     '서울', '사립'],
        ['UNIV_003', '고려대학교',     '서울', '사립'],
        ['UNIV_004', '성균관대학교',   '서울', '사립'],
        ['UNIV_005', '한양대학교',     '서울', '사립'],
        ['UNIV_006', '부산대학교',     '부산', '국립'],
        ['UNIV_007', '경북대학교',     '대구', '국립'],
        ['UNIV_008', '전남대학교',     '광주', '국립'],
        ['UNIV_009', '충남대학교',     '대전', '국립'],
        ['UNIV_010', 'KAIST',          '대전', '국립'],
        ['UNIV_011', '포항공과대학교', '경북', '사립'],
        ['UNIV_012', '인하대학교',     '인천', '사립'],
        ['UNIV_013', '아주대학교',     '경기', '사립'],
        ['UNIV_014', '강원대학교',     '강원', '국립'],
        ['UNIV_015', '제주대학교',     '제주', '국립'],
        ['UNIV_016', '이화여자대학교', '서울', '사립'],
        ['UNIV_017', '숭실대학교',     '서울', '사립'],
        ['UNIV_018', '동아대학교',     '부산', '사립'],
        ['UNIV_019', '전북대학교',     '전북', '국립'],
        ['UNIV_020', '경상국립대학교', '경남', '국립'],
    ],
    [2.5, 5, 2, 2]
)

doc.add_page_break()

# ════════════════════════════════════════════════
# 7. 이후 개발 단계
# ════════════════════════════════════════════════
h1('7. 이후 개발 단계 안내')
body('현재 데모는 Phase 1~2(데이터 수집 + 행동 분석)만 구현되어 있습니다. 이후 단계는 다음과 같이 계획되어 있습니다.')

table(
    ['단계', '기간', '주요 작업'],
    [
        ['Phase 2', '2~3개월', '대학 데이터 Vector DB 임베딩·인덱싱, RAG 검색 API 구현, GPT-Oss 20B 연동'],
        ['Phase 3', '1~2개월', '추천 UI 컴포넌트 개발, 프론트엔드 연동, 피드백 루프(피드백 시그널) 구현'],
        ['Phase 4', '1개월',   'A/B 테스트 실시, 모니터링 대시보드 구축, 성능 최적화'],
    ],
    [2, 2.5, 11]
)

h2('성능 목표 (설계서 9.2 기준)')
table(
    ['지표', '목표값'],
    [
        ['추천 클릭률 (CTR)',     '15% 이상'],
        ['추천 생성 응답 시간',   '2초 이내 (RAG + GPT-Oss 합산)'],
        ['Vector DB 검색 시간',   '100ms 이내 (Top-20 검색)'],
        ['시스템 가용성',         '99.5% 이상'],
    ],
    [7, 8.5]
)

doc.add_paragraph()
note('이 문서는 Phase 1~2 데모 기준으로 작성되었습니다. 이후 단계 구현 완료 시 내용이 업데이트될 예정입니다.')

# ── 저장
output = r'C:\Users\using\HEK\대학알리미_추천시스템_사용자매뉴얼.docx'
doc.save(output)
print(f'저장 완료: {output}')
